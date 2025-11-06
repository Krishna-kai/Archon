"""
Document Processing Utilities

This module provides utilities for extracting text from various document formats
including PDF, Word documents, and plain text files.
"""

import io
import os
from typing import Optional

import httpx

# Removed direct logging import - using unified config

# Import document processing libraries with availability checks
try:
    import PyPDF2

    PYPDF2_AVAILABLE = True
except ImportError:
    PYPDF2_AVAILABLE = False

try:
    import pdfplumber

    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False

try:
    from docx import Document as DocxDocument

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from ..config.logfire_config import get_logger, logfire

logger = get_logger(__name__)


def _preserve_code_blocks_across_pages(text: str) -> str:
    """
    Fix code blocks that were split across PDF page boundaries.
    
    PDFs often break markdown code blocks with page headers like:
    ```python
    def hello():
    --- Page 2 ---
        return "world"
    ```
    
    This function rejoins split code blocks by removing page separators
    that appear within code blocks.
    """
    import re
    
    # Pattern to match page separators that split code blocks
    # Look for: ``` [content] --- Page N --- [content] ```
    page_break_in_code_pattern = r'(```\w*[^\n]*\n(?:[^`]|`(?!``))*)(\n--- Page \d+ ---\n)((?:[^`]|`(?!``))*)```'
    
    # Keep merging until no more splits are found
    while True:
        matches = list(re.finditer(page_break_in_code_pattern, text, re.DOTALL))
        if not matches:
            break
            
        # Replace each match by removing the page separator
        for match in reversed(matches):  # Reverse to maintain positions
            before_page_break = match.group(1)
            page_separator = match.group(2) 
            after_page_break = match.group(3)
            
            # Rejoin the code block without the page separator
            rejoined = f"{before_page_break}\n{after_page_break}```"
            text = text[:match.start()] + rejoined + text[match.end():]
    
    return text


def _clean_html_to_text(html_content: str) -> str:
    """
    Clean HTML tags and convert to plain text suitable for RAG.
    Preserves code blocks and important structure while removing markup.
    """
    import re
    
    # First preserve code blocks with their content before general cleaning
    # This ensures code blocks remain intact for extraction
    code_blocks = []
    
    # Find and temporarily replace code blocks to preserve them
    code_patterns = [
        r'<pre><code[^>]*>(.*?)</code></pre>',
        r'<code[^>]*>(.*?)</code>',
        r'<pre[^>]*>(.*?)</pre>',
    ]
    
    processed_html = html_content
    placeholder_map = {}
    
    for pattern in code_patterns:
        matches = list(re.finditer(pattern, processed_html, re.DOTALL | re.IGNORECASE))
        for i, match in enumerate(reversed(matches)):  # Reverse to maintain positions
            # Extract code content and clean HTML entities
            code_content = match.group(1)
            # Clean HTML entities and span tags from code
            code_content = re.sub(r'<span[^>]*>', '', code_content)
            code_content = re.sub(r'</span>', '', code_content)
            code_content = re.sub(r'&lt;', '<', code_content)
            code_content = re.sub(r'&gt;', '>', code_content)
            code_content = re.sub(r'&amp;', '&', code_content)
            code_content = re.sub(r'&quot;', '"', code_content)
            code_content = re.sub(r'&#39;', "'", code_content)
            
            # Create placeholder
            placeholder = f"__CODE_BLOCK_{len(placeholder_map)}__"
            placeholder_map[placeholder] = code_content.strip()
            
            # Replace in HTML
            processed_html = processed_html[:match.start()] + placeholder + processed_html[match.end():]
    
    # Now clean all remaining HTML tags
    # Remove script and style content entirely
    processed_html = re.sub(r'<script[^>]*>.*?</script>', '', processed_html, flags=re.DOTALL | re.IGNORECASE)
    processed_html = re.sub(r'<style[^>]*>.*?</style>', '', processed_html, flags=re.DOTALL | re.IGNORECASE)
    
    # Convert common HTML elements to readable text
    # Headers
    processed_html = re.sub(r'<h[1-6][^>]*>(.*?)</h[1-6]>', r'\n\n\1\n\n', processed_html, flags=re.DOTALL | re.IGNORECASE)
    # Paragraphs
    processed_html = re.sub(r'<p[^>]*>(.*?)</p>', r'\1\n\n', processed_html, flags=re.DOTALL | re.IGNORECASE)
    # Line breaks
    processed_html = re.sub(r'<br\s*/?>', '\n', processed_html, flags=re.IGNORECASE)
    # List items
    processed_html = re.sub(r'<li[^>]*>(.*?)</li>', r'• \1\n', processed_html, flags=re.DOTALL | re.IGNORECASE)
    
    # Remove all remaining HTML tags
    processed_html = re.sub(r'<[^>]+>', '', processed_html)
    
    # Clean up HTML entities
    processed_html = re.sub(r'&nbsp;', ' ', processed_html)
    processed_html = re.sub(r'&lt;', '<', processed_html)
    processed_html = re.sub(r'&gt;', '>', processed_html)
    processed_html = re.sub(r'&amp;', '&', processed_html)
    processed_html = re.sub(r'&quot;', '"', processed_html)
    processed_html = re.sub(r'&#39;', "'", processed_html)
    processed_html = re.sub(r'&#x27;', "'", processed_html)
    
    # Restore code blocks
    for placeholder, code_content in placeholder_map.items():
        processed_html = processed_html.replace(placeholder, f"\n\n```\n{code_content}\n```\n\n")
    
    # Clean up excessive whitespace
    processed_html = re.sub(r'\n\s*\n\s*\n', '\n\n', processed_html)  # Max 2 consecutive newlines
    processed_html = re.sub(r'[ \t]+', ' ', processed_html)  # Multiple spaces to single space
    
    return processed_html.strip()


def detect_pdf_type(file_content: bytes) -> str:
    """
    Detect if PDF is text-based, scanned, or mixed content.

    This function analyzes a PDF to determine if OCR is needed:
    - 'text_based': PDF contains extractable text, no OCR needed
    - 'scanned': PDF is image-based, requires OCR
    - 'mixed': PDF has both text and images, may benefit from OCR
    - 'unknown': Unable to determine (treat as requiring OCR)

    Args:
        file_content: Raw PDF bytes

    Returns:
        One of: 'text_based', 'scanned', 'mixed', 'unknown'
    """
    try:
        import fitz  # PyMuPDF

        doc = fitz.open(stream=file_content, filetype="pdf")

        if len(doc) == 0:
            doc.close()
            return "unknown"

        # Sample first 3 pages (or all pages if fewer than 3)
        sample_pages = min(3, len(doc))
        total_text_chars = 0
        total_image_count = 0

        for page_num in range(sample_pages):
            page = doc[page_num]

            # Check for extractable text
            text = page.get_text()
            total_text_chars += len(text.strip())

            # Check for images
            images = page.get_images()
            total_image_count += len(images)

        doc.close()

        # Classification logic
        # IMPROVED: Prioritize text content for research papers with figures

        # High text content threshold for academic papers with embedded text
        # Research papers with 8000+ chars (~1500+ words) are clearly text-based
        # even if they contain diagrams, figures, or equations
        if total_text_chars > 8000:
            logger.info(
                f"PDF classified as text_based (heavy text content: "
                f"{total_text_chars} chars, {total_image_count} images)"
            )
            return "text_based"

        # Medium text content - check image ratio
        # Papers with 3000+ chars but few images are still text-based
        elif total_text_chars > 3000:
            if total_image_count > sample_pages * 2:  # More than 2 images per page
                logger.info(
                    f"PDF classified as mixed (moderate text: {total_text_chars} chars, "
                    f"many images: {total_image_count})"
                )
                return "mixed"
            logger.info(f"PDF classified as text_based ({total_text_chars} chars)")
            return "text_based"

        # Low text content - original logic
        elif total_text_chars > 500:
            if total_image_count > sample_pages:  # Also has significant images
                return "mixed"
            return "text_based"

        # If no text but has images
        elif total_image_count > 0:
            return "scanned"
        else:
            # No text and no images (unusual, could be corrupted)
            return "unknown"

    except Exception as e:
        logger.warning(f"PDF type detection failed: {e}")
        return "unknown"


async def extract_text_from_parser_service(
    file_content: bytes,
    filename: str,
    extract_formulas: bool = True,
    extract_tables: bool = True
) -> str:
    """
    Extract text from PDF using Parser Service (Docling + LaTeX-OCR orchestrator).

    This provides the highest quality extraction with:
    - Document structure preservation
    - LaTeX formula extraction
    - Table extraction
    - Markdown formatting

    Args:
        file_content: Raw PDF bytes
        filename: Name of the file
        extract_formulas: Whether to extract LaTeX formulas
        extract_tables: Whether to extract tables

    Returns:
        Extracted markdown content with structure, formulas, and tables

    Raises:
        Exception: If Parser Service is unavailable or extraction fails
    """
    parser_service_url = os.getenv("PARSER_SERVICE_URL", "http://parser-service:9004")

    try:
        async with httpx.AsyncClient(timeout=120.0) as client:
            # Prepare multipart form data
            files = {
                "file": (filename, file_content, "application/pdf")
            }
            data = {
                "extract_formulas": str(extract_formulas).lower(),
                "extract_tables": str(extract_tables).lower()
            }

            logger.info(f"Calling Parser Service for {filename} at {parser_service_url}")

            # Call Parser Service
            response = await client.post(
                f"{parser_service_url}/parse",
                files=files,
                data=data
            )
            response.raise_for_status()

            # Parse response
            result = response.json()

            if not result.get("success"):
                error_msg = result.get("error", "Unknown error")
                raise Exception(f"Parser Service failed: {error_msg}")

            markdown = result.get("markdown", "")
            metadata = result.get("metadata", {})

            # Log extraction statistics
            logger.info(
                f"Parser Service extraction complete for {filename}: "
                f"{metadata.get('pages', 0)} pages, "
                f"{metadata.get('formulas_extracted', 0)} formulas, "
                f"{metadata.get('processing_time_ms', 0)}ms"
            )

            return markdown

    except httpx.ConnectError as e:
        logger.warning(f"Parser Service unavailable at {parser_service_url}: {e}")
        raise Exception(f"Parser Service connection failed: {e}")
    except httpx.TimeoutException as e:
        logger.warning(f"Parser Service timeout for {filename}: {e}")
        raise Exception(f"Parser Service timeout: {e}")
    except httpx.HTTPStatusError as e:
        logger.warning(f"Parser Service HTTP error for {filename}: {e}")
        raise Exception(f"Parser Service HTTP error: {e}")
    except Exception as e:
        logger.warning(f"Parser Service extraction failed for {filename}: {e}")
        raise


async def extract_text_from_mineru(
    file_content: bytes,
    filename: str,
    extract_charts: bool = False,
    chart_provider: str = "auto"
) -> tuple[str, list[dict]]:
    """
    Extract text and images from PDF using MinerU (Magic-PDF) for high-accuracy formula extraction.

    Args:
        file_content: Raw PDF bytes
        filename: Name of the PDF file
        extract_charts: Whether to extract chart data from images
        chart_provider: Chart extraction provider ("auto", "local", "claude")

    Returns:
        Tuple of (markdown_text, image_data_list)
        - markdown_text: Extracted markdown with formulas, tables, and text
        - image_data_list: List of ImageData dicts with base64-encoded images

    Raises:
        Exception: If MinerU extraction fails
    """
    try:
        from ..services.mineru_service import get_mineru_service

        mineru_service = get_mineru_service()

        if not mineru_service.is_available():
            raise Exception("MinerU is not installed")

        logger.info(f"Processing PDF with MinerU: {filename}")

        # Determine device based on service configuration
        # If using HTTP service, it runs natively on Mac with MPS
        # If using local CLI, it runs in Docker with CPU only
        device = "mps" if os.getenv("MINERU_SERVICE_URL") else "cpu"

        # Process PDF with MinerU
        success, result = await mineru_service.process_pdf(
            file_content=file_content,
            filename=filename,
            device=device,
            lang="en",
            extract_charts=extract_charts,
            chart_provider=chart_provider,
        )

        if not success:
            error_msg = result.get("error", "Unknown error")
            raise Exception(f"MinerU extraction failed: {error_msg}")

        markdown = result.get("markdown", "")
        metadata = result.get("metadata", {})
        images = result.get("charts", [])  # "charts" key contains ImageData objects with base64

        # Log extraction statistics
        logger.info(
            f"MinerU extraction complete for {filename}: "
            f"{metadata.get('pages', 0)} pages, "
            f"{metadata.get('formulas_count', 0)} formulas, "
            f"{metadata.get('tables_count', 0)} tables, "
            f"{len(images)} images"
        )

        return markdown, images

    except Exception as e:
        logger.warning(f"MinerU extraction failed for {filename}: {e}")
        raise


async def extract_text_from_document(
    file_content: bytes,
    filename: str,
    content_type: str,
    use_ocr: bool = False,
    use_mineru: bool = False,
    extract_charts: bool = False,
    chart_provider: str = "auto",
) -> tuple[str, list[dict]]:
    """
    Extract text and images from various document formats with intelligent OCR selection.

    This function implements intelligent document analysis to determine the best
    extraction method:
    - For PDFs with formulas: Uses MinerU when use_mineru=True (highest accuracy)
    - For images: Uses DeepSeek OCR when use_ocr=True
    - For text-based PDFs: Skips OCR, uses standard extraction (faster)
    - For scanned PDFs: Uses OCRmyPDF → Tesseract fallback chain
    - For mixed PDFs: Uses OCRmyPDF to handle both text and images
    - For charts/diagrams: Extracts chart data when extract_charts=True

    Args:
        file_content: Raw file bytes
        filename: Name of the file
        content_type: MIME type of the file
        use_ocr: Whether to use OCR for images and scanned PDFs
        use_mineru: Whether to use MinerU for PDF formula extraction
        extract_charts: Whether to extract chart data from images
        chart_provider: Chart extraction provider ("auto", "local", "claude")

    Returns:
        Tuple of (text_content, image_data_list)
        - text_content: Extracted text
        - image_data_list: List of ImageData dicts (empty if not using MinerU)

    Raises:
        ValueError: If the file format is not supported
        Exception: If extraction fails
    """
    try:
        # Image files with OCR
        if use_ocr and (
            content_type.startswith("image/") or
            filename.lower().endswith((".jpg", ".jpeg", ".png", ".bmp", ".tiff", ".tif"))
        ):
            text = await extract_text_from_image_ocr(file_content, filename)
            return text, []  # Images extracted separately, not from OCR text

        # PDF files - intelligent OCR selection based on PDF type
        if content_type == "application/pdf" or filename.lower().endswith(".pdf"):
            # Check file size first
            file_size_mb = len(file_content) / (1024 * 1024)

            # Priority 1: MinerU (highest accuracy for formulas and tables)
            if use_mineru:
                try:
                    logger.info(
                        f"Attempting MinerU extraction for {filename} ({file_size_mb:.2f} MB)"
                    )
                    return await extract_text_from_mineru(file_content, filename, extract_charts, chart_provider)
                except Exception as mineru_error:
                    logger.warning(
                        f"MinerU extraction failed for {filename}, falling back to Parser Service: {mineru_error}"
                    )
                    # Continue to fallback methods below

            # Priority 2: Parser Service (good quality with structure + formulas + tables)
            if use_ocr:
                # Skip for very large files (>100MB) to avoid timeout
                if file_size_mb < 100:
                    try:
                        logger.info(
                            f"Attempting Parser Service extraction for {filename} ({file_size_mb:.2f} MB)"
                        )
                        text = await extract_text_from_parser_service(file_content, filename)
                        return text, []
                    except Exception as parser_error:
                        logger.warning(
                            f"Parser Service unavailable for {filename}, falling back to legacy OCR: {parser_error}"
                        )
                        # Continue to fallback methods below
                else:
                    logger.info(
                        f"PDF {filename} is too large for Parser Service ({file_size_mb:.2f} MB), "
                        f"using legacy OCR"
                    )

                # Fallback: Detect PDF type for legacy OCR method selection
                pdf_type = detect_pdf_type(file_content)
                logger.info(f"PDF type detected for {filename}: {pdf_type}")

                # Text-based PDFs don't need OCR
                if pdf_type == "text_based":
                    logger.info(f"PDF {filename} has extractable text, skipping OCR for performance")
                    text = extract_text_from_pdf(file_content)
                    return text, []

                # Scanned PDFs need OCR - choose method based on file size
                elif pdf_type == "scanned":
                    logger.info(f"Scanned PDF detected: {filename} ({file_size_mb:.2f} MB)")

                    if file_size_mb < 5:
                        # Small files: use fast OCRmyPDF
                        try:
                            text = await extract_text_from_pdf_ocr(file_content, filename)
                            return text, []
                        except Exception as ocr_error:
                            logger.warning(
                                f"PDF OCR failed for {filename}, falling back to standard extraction: {ocr_error}"
                            )
                            text = extract_text_from_pdf(file_content)
                            return text, []
                    else:
                        # Large files: use Tesseract directly (more memory efficient)
                        try:
                            logger.info(f"Large scanned PDF, using Tesseract directly for {filename}")
                            text = await extract_text_from_pdf_tesseract(file_content, filename)
                            return text, []
                        except Exception as tesseract_error:
                            logger.warning(f"Tesseract failed for {filename}, trying OCRmyPDF: {tesseract_error}")
                            try:
                                text = await extract_text_from_pdf_ocrmypdf(file_content, filename)
                                return text, []
                            except Exception:
                                text = extract_text_from_pdf(file_content)
                                return text, []

                # Mixed or unknown: use default fallback chain (OCRmyPDF → Tesseract)
                else:
                    try:
                        text = await extract_text_from_pdf_ocr(file_content, filename)
                        return text, []
                    except Exception as ocr_error:
                        logger.warning(
                            f"PDF OCR failed for {filename}, falling back to standard extraction: {ocr_error}"
                        )
                        text = extract_text_from_pdf(file_content)
                        return text, []

            # use_ocr=false: standard extraction without OCR
            text = extract_text_from_pdf(file_content)
            return text, []

        # Word documents
        elif content_type in [
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "application/msword",
        ] or filename.lower().endswith((".docx", ".doc")):
            text = extract_text_from_docx(file_content)
            return text, []

        # HTML files - clean tags and extract text
        elif content_type == "text/html" or filename.lower().endswith((".html", ".htm")):
            # Decode HTML and clean tags for RAG
            html_text = file_content.decode("utf-8", errors="ignore").strip()
            if not html_text:
                raise ValueError(f"The file {filename} appears to be empty.")
            return _clean_html_to_text(html_text)

        # Text files (markdown, txt, etc.)
        elif content_type.startswith("text/") or filename.lower().endswith((
            ".txt",
            ".md",
            ".markdown",
            ".rst",
        )):
            # Decode text and check if it has content
            text = file_content.decode("utf-8", errors="ignore").strip()
            if not text:
                raise ValueError(f"The file {filename} appears to be empty.")
            return text

        else:
            raise ValueError(f"Unsupported file format: {content_type} ({filename})")

    except ValueError:
        # Re-raise ValueError with original message for unsupported formats
        raise
    except Exception as e:
        logfire.error(
            "Document text extraction failed",
            filename=filename,
            content_type=content_type,
            error=str(e),
        )
        # Re-raise with context, preserving original exception chain
        raise Exception(f"Failed to extract text from {filename}") from e


def extract_text_from_pdf(file_content: bytes, preserve_layout: bool = True) -> str:
    """
    Extract text from PDF using PyMuPDF with layout preservation.
    Optimized for research papers with multi-column layouts.

    Args:
        file_content: Raw PDF bytes
        preserve_layout: If True, uses advanced flags for better spacing and layout preservation

    Returns:
        Extracted text content

    Raises:
        Exception: If extraction fails
    """
    try:
        import fitz  # PyMuPDF

        doc = fitz.open(stream=file_content, filetype="pdf")
        text_parts = []

        for page_num in range(len(doc)):
            page = doc[page_num]

            if preserve_layout:
                # Use advanced flags for better layout preservation
                # TEXT_PRESERVE_WHITESPACE: Maintains spacing between columns
                # TEXT_PRESERVE_IMAGES: Marks where images are (for reference)
                try:
                    text = page.get_text(
                        "text",
                        flags=fitz.TEXT_PRESERVE_WHITESPACE |
                              fitz.TEXT_PRESERVE_IMAGES
                    )
                except Exception as e:
                    # Fallback to default extraction if flags fail
                    logger.warning(f"PyMuPDF advanced extraction failed on page {page_num + 1}: {e}, using default")
                    text = page.get_text()
            else:
                text = page.get_text()

            if text and text.strip():
                text_parts.append(text)

        doc.close()

        if not text_parts:
            raise ValueError(
                "No text extracted from PDF: file may be empty, images-only, "
                "or scanned document without OCR"
            )

        extracted = "\n\n".join(text_parts)

        # Log extraction quality metrics
        word_count = len(extracted.split())
        logger.info(
            f"PyMuPDF extracted {len(extracted)} chars, {word_count} words "
            f"from {len(text_parts)} pages (preserve_layout={preserve_layout})"
        )

        return extracted

    except Exception as e:
        logger.error(f"PyMuPDF extraction failed: {e}")
        raise Exception(f"Failed to extract text from PDF using PyMuPDF") from e


def extract_text_from_docx(file_content: bytes) -> str:
    """
    Extract text from Word documents (.docx).

    Args:
        file_content: Raw DOCX bytes

    Returns:
        Extracted text content
    """
    if not DOCX_AVAILABLE:
        raise Exception("python-docx library not available. Please install python-docx.")

    try:
        doc = DocxDocument(io.BytesIO(file_content))
        text_content = []

        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text)

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_content.append(" | ".join(row_text))

        if not text_content:
            raise ValueError("No text content found in document")

        return "\n\n".join(text_content)

    except Exception as e:
        raise Exception("Failed to extract text from Word document") from e


async def extract_text_from_image_ocr(file_content: bytes, filename: str) -> str:
    """
    Extract text from image files using OCR.

    Args:
        file_content: Raw image bytes
        filename: Name of the file

    Returns:
        Extracted text content

    Raises:
        Exception: If OCR fails or service is unavailable
    """
    from ..services.ocr_service import get_ocr_service

    logger.info(f"Using OCR to extract text from image: {filename}")

    # Get OCR service instance
    ocr_service = get_ocr_service()

    # Check OCR service health
    try:
        health = await ocr_service.check_health()
        if health.get("status") != "healthy":
            raise Exception(
                f"OCR service is not available: {health.get('error', 'Unknown error')}"
            )
    except Exception as e:
        logger.error(f"OCR service health check failed: {e}")
        raise Exception(
            "OCR service is not available. Please ensure DeepSeek OCR service is running."
        ) from e

    # Perform OCR
    try:
        success, result = await ocr_service.ocr_image(file_content, filename)

        if success:
            extracted_text = result.get('text', '')
            if not extracted_text or not extracted_text.strip():
                raise ValueError(f"OCR extracted no text from {filename}")

            logger.info(
                f"OCR extraction successful | filename={filename} | "
                f"length={len(extracted_text)} | backend={result.get('backend')}"
            )
            return extracted_text
        else:
            error_msg = result.get('error', 'OCR failed')
            raise Exception(f"OCR failed: {error_msg}")

    except Exception as e:
        logger.error(f"OCR extraction failed for {filename}: {e}")
        raise Exception(f"Failed to extract text from image using OCR") from e


async def extract_text_from_pdf_tesseract(file_content: bytes, filename: str) -> str:
    """
    Extract text from PDF files using Tesseract OCR with optimized settings.

    Uses Tesseract Docker container for OCR processing. Optimized for research papers
    with improved PSM (Page Segmentation Mode) and OEM (OCR Engine Mode) settings.

    Args:
        file_content: Raw PDF bytes
        filename: Name of the file

    Returns:
        Extracted text content

    Raises:
        Exception: If OCR fails or Docker is unavailable
    """
    import subprocess
    import tempfile
    import os
    import time
    import fitz  # PyMuPDF

    logger.info(f"Using Tesseract OCR (improved) for PDF: {filename}")
    start_time = time.time()

    # Convert PDF pages to images, then OCR each page
    try:
        # Write PDF to temp file
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as pdf_temp:
            pdf_temp.write(file_content)
            pdf_path = pdf_temp.name

        # Open PDF with PyMuPDF to extract pages as images
        doc = fitz.open(pdf_path)
        text_parts = []

        with tempfile.TemporaryDirectory() as temp_dir:
            for page_num in range(len(doc)):
                page = doc[page_num]

                # Render page to high-resolution image for better OCR
                pix = page.get_pixmap(dpi=300)  # 300 DPI is standard for OCR
                img_path = os.path.join(temp_dir, f"page_{page_num}.png")
                pix.save(img_path)

                # Optimized Tesseract command for research papers
                cmd = [
                    "docker", "run", "--rm", "-i",
                    "jitesoft/tesseract-ocr",
                    "tesseract",
                    "stdin",  # Read from stdin
                    "stdout",  # Write to stdout
                    "-l", "eng",  # English language
                    "--psm", "1",  # Auto page segmentation with OSD (best for papers)
                    "--oem", "3",  # Use LSTM neural network engine
                    "-c", "preserve_interword_spaces=1"  # Better word spacing
                ]

                # Read image and send to Tesseract
                with open(img_path, "rb") as img_file:
                    img_data = img_file.read()

                result = subprocess.run(
                    cmd,
                    input=img_data,
                    capture_output=True,
                    timeout=180  # 3 minute timeout per page
                )

                if result.returncode == 0 and result.stdout:
                    # Decode bytes to string
                    text = result.stdout.decode('utf-8', errors='ignore')
                    if text.strip():  # Only add non-empty text
                        text_parts.append(f"--- Page {page_num + 1} ---\n{text}")
                        logger.debug(f"Tesseract page {page_num + 1}: extracted {len(text)} chars")
                else:
                    error_msg = result.stderr.decode('utf-8', errors='ignore') if result.stderr else "Unknown error"
                    logger.warning(f"Tesseract failed on page {page_num + 1}: {error_msg}")

        doc.close()
        os.remove(pdf_path)

        elapsed = time.time() - start_time
        extracted_text = "\n\n".join(text_parts)

        if not extracted_text or not extracted_text.strip():
            raise ValueError(f"Tesseract extracted no text from {filename}")

        word_count = len(extracted_text.split())
        logger.info(
            f"Tesseract extraction successful | filename={filename} | "
            f"pages={len(text_parts)} | chars={len(extracted_text)} | "
            f"words={word_count} | time={elapsed:.2f}s | backend=tesseract"
        )

        return extracted_text

    except subprocess.TimeoutExpired:
        raise Exception(f"Tesseract OCR timed out after {time.time() - start_time:.1f}s")
    except Exception as e:
        logger.error(f"Tesseract OCR failed: {e}")
        raise Exception(f"Failed to run Tesseract OCR: {e}")


async def extract_text_from_pdf_ocrmypdf(
    file_content: bytes,
    filename: str,
    quality_mode: str = "balanced"
) -> str:
    """
    Extract text from PDF files using OCRmyPDF with stdin/stdout pattern.

    Uses OCRmyPDF Docker container for OCR processing. Ideal for scanned PDFs
    or image-heavy documents. Optimized for Apple Silicon (M1/M2/M3/M4).

    Uses stdin/stdout to avoid Docker-in-Docker volume mounting issues.

    Args:
        file_content: Raw PDF bytes
        filename: Name of the file
        quality_mode: "fast", "balanced", or "quality" (default: "balanced")

    Returns:
        Extracted text content

    Raises:
        Exception: If OCR fails or Docker is unavailable
    """
    import subprocess
    import tempfile
    import os
    import time
    import fitz  # PyMuPDF

    logger.info(f"Using OCRmyPDF ({quality_mode} mode) for PDF: {filename}")
    start_time = time.time()

    # Run OCRmyPDF with stdin/stdout to avoid volume mounting issues
    try:
        # Quality mode configurations
        if quality_mode == "fast":
            cmd = [
                "docker", "run", "--rm", "-i",
                "jbarlow83/ocrmypdf-alpine",
                "--deskew",
                "--optimize", "1",
                "--skip-text",  # Skip pages that already have text
                "-", "-"  # stdin to stdout
            ]
            timeout = 300
        elif quality_mode == "quality":
            cmd = [
                "docker", "run", "--rm", "-i",
                "jbarlow83/ocrmypdf-alpine",
                "--deskew",
                "--rotate-pages",
                "--clean",
                "--optimize", "3",
                "--skip-text",
                "--tesseract-pagesegmode", "1",
                "-", "-"
            ]
            timeout = 600
        else:  # balanced
            cmd = [
                "docker", "run", "--rm", "-i",
                "jbarlow83/ocrmypdf-alpine",
                "--deskew",
                "--rotate-pages",
                "--optimize", "2",
                "--skip-text",
                "-", "-"
            ]
            timeout = 400

        logger.info(f"Running OCRmyPDF on {filename} ({len(file_content)} bytes)")

        result = subprocess.run(
            cmd,
            input=file_content,
            capture_output=True,
            timeout=timeout
        )

        elapsed = time.time() - start_time

        if result.returncode != 0:
            error_msg = result.stderr.decode("utf-8", errors="replace") if result.stderr else "Unknown error"
            logger.error(f"OCRmyPDF failed after {elapsed:.1f}s: {error_msg}")
            raise Exception(f"OCRmyPDF processing failed: {error_msg}")

        logger.info(f"OCRmyPDF completed in {elapsed:.2f}s ({len(result.stdout)} bytes)")

        # OCR'd PDF is in result.stdout
        ocr_pdf_bytes = result.stdout

        if not ocr_pdf_bytes:
            raise Exception("OCRmyPDF returned empty output")

    except subprocess.TimeoutExpired:
        elapsed = time.time() - start_time
        raise Exception(f"OCRmyPDF timed out after {elapsed:.1f}s")
    except Exception as e:
        raise Exception(f"Failed to run OCRmyPDF: {e}")

    # Extract text from OCR'd PDF using PyMuPDF
    try:
        # Open PDF from bytes directly (no temp file needed)
        doc = fitz.open(stream=ocr_pdf_bytes, filetype="pdf")
        text_parts = []

        for page_num in range(len(doc)):
            page = doc[page_num]
            text = page.get_text()
            text_parts.append(f"\n--- Page {page_num + 1} ---\n{text}")

        doc.close()
        extracted_text = "\n".join(text_parts)

        if not extracted_text or not extracted_text.strip():
            raise ValueError(f"OCR extracted no text from {filename}")

        logger.info(
            f"OCR extraction successful | filename={filename} | "
            f"pages={len(text_parts)} | chars={len(extracted_text)} | backend=ocrmypdf"
        )

        return extracted_text

    except Exception as e:
        raise Exception(f"Failed to extract text from OCR'd PDF: {e}")


async def extract_text_from_pdf_ocr(file_content: bytes, filename: str) -> str:
    """
    Extract text from PDF files using OCR with intelligent fallback chain.

    Fallback chain:
    1. OCRmyPDF (primary) - ARM64-native, best for academic PDFs
    2. Tesseract (backup) - More configurable, custom processing
    3. If both fail, exception is raised (caller handles standard extraction fallback)

    Args:
        file_content: Raw PDF bytes
        filename: Name of the file

    Returns:
        Extracted text content

    Raises:
        Exception: If all OCR methods fail
    """
    # Try OCRmyPDF first (primary, fastest, best quality)
    try:
        logger.info(f"Attempting OCR extraction with OCRmyPDF (primary) for {filename}")
        return await extract_text_from_pdf_ocrmypdf(file_content, filename)
    except Exception as ocrmypdf_error:
        logger.warning(f"OCRmyPDF failed for {filename}: {ocrmypdf_error}")

        # Fallback to Tesseract
        try:
            logger.info(f"Falling back to Tesseract OCR for {filename}")
            return await extract_text_from_pdf_tesseract(file_content, filename)
        except Exception as tesseract_error:
            logger.error(f"Tesseract OCR also failed for {filename}: {tesseract_error}")

            # Both OCR methods failed, raise combined error
            raise Exception(
                f"All OCR methods failed. "
                f"OCRmyPDF: {str(ocrmypdf_error)[:100]}... "
                f"Tesseract: {str(tesseract_error)[:100]}..."
            )
